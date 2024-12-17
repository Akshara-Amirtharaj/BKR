from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import platform
import subprocess
import base64

app = Flask(__name__)

# Path to the text file for storing base number and counter
SERIAL_FILE = "serial_data.txt"

def get_serial_number():
    # Read base number and counter from file
    try:
        with open(SERIAL_FILE, "r") as f:
            base_number, counter = map(int, f.read().strip().split(","))
    except FileNotFoundError:
        base_number, counter = 701, 0
        with open(SERIAL_FILE, "w") as f:
            f.write(f"{base_number},{counter}")

    # Calculate current serial number
    serial_number = base_number + counter

    # Increment the counter and update the file
    with open(SERIAL_FILE, "w") as f:
        f.write(f"{base_number},{counter + 1}")

    return serial_number

def generate_reference_number(company_name="BKR"):
    """
    Generate the full reference number in the format: BKRMM-YYYY-CR<serial>.
    """
    current_month = datetime.now().strftime("%m")
    current_year = datetime.now().strftime("%Y")
    serial_number = get_serial_number()
    return f"{company_name}{current_month}-{current_year}-CR{serial_number}"

def replace_placeholders(doc, placeholders):
    """Replace placeholders in a Word document, including paragraphs and tables."""
    
    def replace_in_paragraph(paragraph, key, value):
        """Replace placeholders in a single paragraph, handling split runs."""
        full_text = "".join(run.text for run in paragraph.runs)
        if key in full_text:
            full_text = full_text.replace(key, value)
            for run in paragraph.runs:
                run.text = ""  # Clear all runs
            paragraph.runs[0].text = full_text  # Add the replaced text back

    def replace_in_cell(cell, placeholders):
        """Replace placeholders inside a table cell."""
        for para in cell.paragraphs:
            for key, value in placeholders.items():
                replace_in_paragraph(para, key, value)

    # Replace placeholders in all paragraphs
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            replace_in_paragraph(para, key, value)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell, placeholders)

    return doc

def replace_placeholders_vat(doc, placeholders):
    """Replace placeholders in a Word document, maintaining original formatting."""
    
    def replace_in_paragraph(paragraph, key, value):
        """Replace placeholders in a paragraph, preserving formatting."""
        for run in paragraph.runs:
            if key in run.text:
                # Replace placeholder text
                run.text = run.text.replace(key, value)
                # Retain original font style and size
                run.font.name = paragraph.style.font.name
                run.font.size = paragraph.style.font.size

    def replace_in_cell(cell, placeholders):
        """Replace placeholders inside a table cell."""
        for para in cell.paragraphs:
            for key, value in placeholders.items():
                replace_in_paragraph(para, key, value)

    # Replace placeholders in all paragraphs
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            replace_in_paragraph(para, key, value)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell, placeholders)

    return doc

def apply_image_placeholder(doc, placeholder_key, image_file):
    """Replace a placeholder with an image in the Word document, resized to fit a specific size."""
    try:
        # Iterate through tables to find the placeholder
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if placeholder_key in para.text:
                            # Clear the placeholder text
                            para.text = ""
                            run = para.add_run()
                            # Resize the image to a fixed width and height (e.g., 1.5 x 1.5 inches)
                            run.add_picture(image_file, width=Inches(1.5), height=Inches(0.75))  # Adjust dimensions
                            return doc  # Exit after placing the image

        # Check for the placeholder in paragraphs outside tables
        for para in doc.paragraphs:
            if placeholder_key in para.text:
                # Clear the placeholder text
                para.text = ""
                run = para.add_run()
                # Resize the image to a fixed width and height (e.g., 1.5 x 1.5 inches)
                run.add_picture(image_file, width=Inches(1.5), height=Inches(0.75))  # Adjust dimensions
                return doc  # Exit after placing the image

        raise ValueError(f"Placeholder '{placeholder_key}' not found in the document.")
    except Exception as e:
        raise Exception(f"Error inserting image: {e}")

def convert_to_pdf(doc_path, pdf_path):
    doc_path = os.path.abspath(doc_path)
    pdf_path = os.path.abspath(pdf_path)

    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Word document not found at {doc_path}")

    if platform.system() == "Windows":
        try:
            import comtypes.client
            import pythoncom
            pythoncom.CoInitialize()
            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(doc_path)
            doc.SaveAs(pdf_path, FileFormat=17)
            doc.Close()
            word.Quit()
        except Exception as e:
            raise Exception(f"Error using COM on Windows: {e}")
    else:
        try:
            subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), doc_path],
                check=True
            )
        except subprocess.CalledProcessError as e:
            raise Exception(f"Error using LibreOffice: {e}")

@app.route('/generate-document', methods=['POST'])
def generate_document():
    data = request.json

    # Parse inputs
    template_type = data.get("template_type")
    placeholders = data.get("placeholders", {})
    signature_image_base64 = data.get("signature_image")  # Base64 encoded image

    try:
        # Generate Reference Number
        reference_number = generate_reference_number()
        placeholders["<<Reference Number>>"] = reference_number

        # Load correct template
        if template_type == "VAT":
            template_path = "SAMPLE VAT registration and VAT filling -SME package.docx"
            doc = Document(template_path)

            # Replace placeholders
            doc = replace_placeholders_vat(doc, placeholders)
        elif template_type == "Service Agreement":
            template_path = "SAMPLE Service Agreement -Company formation -Bahrain - Filled.docx"
            doc = Document(template_path)

            # Replace placeholders
            doc = replace_placeholders(doc, placeholders)
        else:
            return jsonify({"status": "error", "message": "Invalid template type!"}), 400

        doc = Document(template_path)

        # Replace placeholders
        doc = replace_placeholders(doc, placeholders)

        # Handle signature image if provided
        if signature_image_base64:
            image_path = "signature_image.png"
            with open(image_path, "wb") as img_file:
                img_file.write(base64.b64decode(signature_image_base64))
            doc = apply_image_placeholder(doc, "<<Signature Image>>", image_path)

        # Save document
        word_output = f"{template_type} {reference_number}.docx"
        pdf_output = word_output.replace(".docx", ".pdf")

        doc.save(word_output)
        convert_to_pdf(word_output, pdf_output)

        # Return response
        return jsonify({
            "status": "success",
            "reference_number": reference_number,
            "word_document": word_output,
            "pdf_document": pdf_output
        })

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    if os.path.exists(filename):
        return send_file(filename, as_attachment=True)
    else:
        return jsonify({"status": "error", "message": "File not found!"}), 404
@app.route("/", methods=["GET"])
def home():
    return jsonify({
        "message": "Document Generation API is running!",
        "endpoints": {
            "POST /generate-document": "Generate a document with placeholders",
            "GET /download/<filename>": "Download a generated document"
        }
    })

if __name__ == '__main__':
    import os

    port = int(os.environ.get("PORT", 8080))  # Default to 8501 if PORT is not set
    app.run(host="0.0.0.0", port=port, debug=True)

    
